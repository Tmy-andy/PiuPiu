# print("✅ Flask khởi động...")
import os
import traceback
# print("✅ Flask đang được yêu cầu chạy ở cổng :", os.environ.get("PORT"))
# print("📦 Environment:", dict(os.environ))
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, Response, abort
from docx import Document
from flask_login import login_required, LoginManager, login_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import timedelta
from datetime import datetime
from database import init_app, db
from models import User, MemberID, PointLog, Rule, CharacterAbility, BlacklistEntry, KimBaiLog, PlayerOffRequest, GamePlayer, GameHistory, ActivityLog
from functools import wraps
from sqlalchemy import func, union_all
from sqlalchemy.orm import aliased
from sqlalchemy.inspection import inspect
from flask_sqlalchemy import SQLAlchemy
import logging
from zipfile import ZipFile
import tempfile
import csv
import io
from flask_migrate import Migrate
from dotenv import load_dotenv
from flask import jsonify
import pytz
from flask_wtf.csrf import CSRFProtect

load_dotenv()

def get_app_version():
    env_version = os.getenv("APP_VERSION")
    if env_version:
        return env_version.strip()
    try:
        with open(os.path.join(os.path.dirname(__file__), "version.txt"), "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        return "v0.0"

APP_VERSION = get_app_version()

def get_app_changelog():
    try:
        with open(os.path.join(os.path.dirname(__file__), "changelog.txt"), "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        return "Không có ghi chú thay đổi."

APP_CHANGELOG = get_app_changelog()

try:
    app = Flask(__name__)
    app.logger.info(f"Ứng dụng khởi động với phiên bản: {APP_VERSION}")
    app.logger.setLevel(logging.DEBUG)
    secret = os.environ.get('SECRET_KEY')
    if not secret:
        raise RuntimeError("SECRET_KEY không được để trống trong production!")
    app.secret_key = secret
    app.permanent_session_lifetime = timedelta(days=30)


    init_app(app)

    with app.app_context():
        db.create_all()

    # print("✅ Flask khởi động...")

except Exception as e:
    print("🛑 Lỗi khi khởi tạo Flask app:")
    traceback.print_exc()

#Session Cookie Security
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SECURE=False,  # Bật nếu dùng HTTPS
    SESSION_COOKIE_SAMESITE='Lax'
)

# Lấy SECRET_KEY từ biến môi trường (Railway Variables)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'fallback_key_dev')
# Bật bảo vệ CSRF
csrf = CSRFProtect(app)

# Cấu hình SQLAlchemy
from flask_cors import CORS
CORS(app, resources={r"/*": {"origins": "https://piupiu-production.up.railway.app/"}})


from models import db

migrate = Migrate(app, db)

# Tạo các bảng nếu chưa có
with app.app_context():
    db.create_all()

from flask_compress import Compress
Compress(app)

@app.after_request
def add_cache_control(response):
    # Nếu là file tĩnh: cache lâu hơn
    if request.path.startswith('/static/'):
        response.headers['Cache-Control'] = 'public, max-age=31536000'  # 1 năm
    else:
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))

        user = User.query.get(session['user_id'])
        if not user or user.role != 'admin':
            flash('Bạn không có quyền truy cập trang này.', 'error')
            return redirect(url_for('dashboard'))

        return f(*args, **kwargs)
    return decorated_function

login_manager = LoginManager()
login_manager.init_app(app)
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def log_activity(action, detail=""):
    from models import ActivityLog, User, db
    from flask import session

    user_id = session.get("user_id")
    if user_id:
        vietnam_tz = pytz.timezone('Asia/Ho_Chi_Minh')
        log = ActivityLog(
            user_id=user_id,
            action=action,
            detail=detail,
            timestamp=datetime.now(vietnam_tz)
        )
        db.session.add(log)
        db.session.commit()
        cache.delete('active_logs')  # Xóa cache để reload log mới

from flask_wtf.csrf import generate_csrf

@app.context_processor
def inject_user():
    user_id = session.get('user_id')
    user = User.query.get(user_id) if user_id else None

    warning_count = cache.get("warning_count")
    vietnam_now = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))
    if warning_count is None:
        warning_count = 0
        now = vietnam_now
        users = User.query.all()

        for u in users:
            on_leave = PlayerOffRequest.query.filter(
                PlayerOffRequest.user_id == u.id,
                PlayerOffRequest.start_date <= now.date(),
                PlayerOffRequest.end_date >= now.date()
            ).first()
            if on_leave:
                continue

            last_game = (
                GamePlayer.query
                .filter_by(player_id=u.id)
                .join(GameHistory)
                .order_by(GameHistory.created_at.desc())
                .first()
            )
            last_play_time = last_game.game.created_at if last_game else None
            if last_play_time:
                if last_play_time.tzinfo is None:
                    last_play_time = pytz.utc.localize(last_play_time)
                last_play_time = last_play_time.astimezone(pytz.timezone('Asia/Ho_Chi_Minh'))

            if not last_play_time or (now - last_play_time).days > 7:
                warning_count += 1

        cache.set("warning_count", warning_count, timeout=300)

    effective_theme = get_theme_with_cache(user_id) if user_id else 'default'
    return dict(
        user=user,
        warning_count=warning_count,
        effective_theme=effective_theme,
        csrf_token=generate_csrf  # Thêm dòng này
    )

def get_theme_with_cache(user_id):
    cache_key = f"user_theme:{user_id}"

    # 1. Check cache
    theme = cache.get(cache_key)
    if theme:
        return theme

    # 2. Nếu chưa có: query DB
    user = User.query.get(user_id)
    if not user:
        return 'default'

    # 3. Ưu tiên ngày lễ
    today = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))
    if today.month == 10 and today.day >= 25:
        theme = 'halloween'
    elif today.month == 12 and today.day >= 24:
        theme = 'christmas'
    elif today.month == 1 and today.day <= 2:
        theme = 'newyear'
    else:
        theme = user.theme or 'default'

    # 4. Cache lại 1 tiếng
    cache.set(cache_key, theme, timeout=3600)

    return theme

# Cache
from flask_caching import Cache

cache = Cache(config={
    'CACHE_TYPE': 'SimpleCache',
    'CACHE_DEFAULT_TIMEOUT': 300
})
cache.init_app(app)

# Reset cache nếu có deploy mới
LAST_DEPLOY_VERSION_KEY = "last_deploy_version"

def reset_cache_if_new_version():
    last_version = cache.get(LAST_DEPLOY_VERSION_KEY)
    if last_version != APP_VERSION:
        cache.clear()
        cache.set(LAST_DEPLOY_VERSION_KEY, APP_VERSION)
        app.logger.info(f"🚀 Deploy mới: Cache đã reset! Phiên bản: {APP_VERSION}")

        try:
            with app.app_context():
                admin_user = User.query.filter_by(member_id='ADMIN-001').first()
                if admin_user:
                    detail_msg = f"🚀 Admin đã nâng cấp website lên phiên bản {APP_VERSION}: {APP_CHANGELOG}"
                    log = ActivityLog(
                        user_id=admin_user.id,
                        action="Nâng cấp hệ thống",
                        detail=detail_msg,
                        timestamp=datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))
                    )
                    db.session.add(log)
                    db.session.commit()
        except Exception as e:
            app.logger.error(f"Lỗi ghi ActivityLog deploy: {e}")

# Kiểm tra version mới và reset cache nếu cần
with app.app_context():
    reset_cache_if_new_version()

# Timezone conversion
vietnam_tz = pytz.timezone('Asia/Ho_Chi_Minh')

def convert_logs_timezone(logs, tz=vietnam_tz):
    converted = []
    for log in logs:
        ts = log.timestamp
        if ts.tzinfo is None:
            # Giả định ts đang ở UTC
            ts = pytz.utc.localize(ts)
        # Chuyển sang timezone VN
        ts_local = ts.astimezone(tz)
        # Tạo bản dict mới để giữ nguyên log gốc
        converted.append({
            'timestamp': ts_local,
            'action': log.action,
            'detail': log.detail,
        })
    return converted


# Error handlers
@app.errorhandler(400)
def bad_request(e):
    return render_template('errors/400.html'), 400

@app.errorhandler(403)
def forbidden(e):
    return render_template('errors/403.html'), 403

@app.errorhandler(404)
def not_found(e):
    return render_template("errors/404.html"), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template("errors/500.html"), 500

@app.errorhandler(401)
def handle_unauthorized(error):
    return render_template("errors/unauthorized.html"), error.code

# Routes
@app.route('/ping')
def ping():
    return "pong"

@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        member_id = request.form['member_id']
        password = request.form['password']

        user = User.query.filter_by(member_id=member_id).first()

        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            session.permanent = True
            session['user_id'] = user.id
            session['user_role'] = user.role
            session['display_name'] = user.display_name
            session['theme'] = get_theme(user)
            flash(f'Chào mừng {user.display_name}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Mã thành viên hoặc mật khẩu không đúng.', 'error')

    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        member_id = request.form['member_id']
        display_name = request.form['display_name']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            flash('Mật khẩu xác nhận không khớp.', 'error')
            return render_template('register.html')

        member_id_record = MemberID.query.filter_by(member_id=member_id).first()
        if not member_id_record:
            flash('Mã thành viên không tồn tại.', 'error')
            return render_template('register.html')

        if member_id_record.is_used:
            flash('Mã thành viên đã được sử dụng.', 'error')
            return render_template('register.html')

        existing_user = User.query.filter_by(member_id=member_id).first()
        if existing_user:
            flash('Mã thành viên đã được đăng ký.', 'error')
            return render_template('register.html')

        # Tạo user mới
        password_hash = generate_password_hash(password)
        new_user = User(
            member_id=member_id,
            display_name=display_name,
            password_hash=password_hash,
            role='member',
            points=10
        )
        db.session.add(new_user)
        db.session.commit()

        # Đánh dấu mã thành viên là đã dùng
        member_id_record.is_used = True
        member_id_record.used_by = new_user.id
        db.session.commit()
        log_activity("Đăng ký", f"Người dùng mới {new_user.display_name} (ID {new_user.id}, mã {new_user.member_id}) đã đăng ký.")
        flash('Đăng ký thành công! Vui lòng đăng nhập.', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/dashboard')
@login_required
def dashboard():
    user = User.query.get(session['user_id'])

    if user.role == 'admin':
        total_members = User.query.filter_by(role='member').count()
        total_points = db.session.query(db.func.sum(User.points)).filter_by(role='member').scalar() or 0
        avg_points = total_points / total_members if total_members > 0 else 0

        assigned_members = User.query.filter_by(role='member', assigned_admin_id=user.id).order_by(User.points.desc()).all()

        return render_template(
            'admin_dashboard.html',
            total_members=total_members,
            total_points=total_points,
            avg_points=round(avg_points, 1),
            assigned_members=assigned_members,
            admin_points=user.points  # 👈 thêm dòng này
        )
    else:
        point_logs = db.session.query(PointLog).join(User, PointLog.admin_id == User.id) \
            .filter(PointLog.member_id == user.id) \
            .order_by(PointLog.created_at.desc()).all()

        return render_template('member_dashboard.html', user=user, point_logs=point_logs)

@app.route('/members')
@admin_required
def members():

    Admin = aliased(User)

    # ⚙️ Cấu hình phân trang
    per_page = 20
    page = int(request.args.get('page', 1))
    offset = (page - 1) * per_page

    # 🔎 Tổng số thành viên
    total = User.query.filter_by(role='member').count()
    total_pages = ceil(total / per_page)

    # ⚡ Truy vấn có phân trang + join admin
    results = db.session.query(
        User,
        Admin.display_name.label("admin_name")
    ).outerjoin(Admin, User.assigned_admin_id == Admin.id) \
     .filter(User.role == 'member') \
     .order_by(User.member_id.asc()) \
     .offset(offset).limit(per_page).all()

    # ✅ Gắn admin_name vào user
    members = []
    for user, admin_name in results:
        user.admin_name = admin_name
        members.append(user)

    # ✅ Lấy danh sách admin
    all_admins = User.query.filter_by(role='admin').order_by(User.display_name).all()

    return render_template(
        'members.html',
        members=members,
        all_admins=all_admins,
        total=total,
        page=page,
        total_pages=total_pages
    )

@app.route('/assign_member/<int:user_id>', methods=['POST'])
@admin_required
def assign_member(user_id):
    try:
        new_admin_id = request.form.get('admin_id')

        user = User.query.get(user_id) if user_id else None
        if not user or user.role != 'member':
            flash('Không tìm thấy thành viên hợp lệ.', 'danger')
            return redirect(url_for('members'))

        if new_admin_id:
            try:
                new_admin_id = int(new_admin_id)
                new_admin = User.query.get(new_admin_id)
            except ValueError:
                flash('ID admin không hợp lệ.', 'danger')
                return redirect(url_for('members'))

            if not new_admin or new_admin.role != 'admin':
                flash('Admin không hợp lệ.', 'danger')
                return redirect(url_for('members'))

            user.assigned_admin_id = new_admin.id
        else:
            user.assigned_admin_id = None

        db.session.commit()
        flash(f'Đã cập nhật admin phụ trách cho {user.display_name}.', 'success')
        return redirect(url_for('members'))

    except Exception as e:
        print("Lỗi ở /assign_member:", e)
        flash('Đã xảy ra lỗi nội bộ.', 'danger')
        return redirect(url_for('members'))

@app.route('/member_ids')
@admin_required
def member_ids():
    UsedBy = aliased(User)

    per_page = 30
    page = int(request.args.get('page', 1))
    offset = (page - 1) * per_page

    total = MemberID.query.count()
    total_pages = ceil(total / per_page)

    member_ids = db.session.query(MemberID, UsedBy.display_name.label("used_by_name")) \
        .outerjoin(UsedBy, MemberID.used_by == UsedBy.id) \
        .order_by(MemberID.member_id.asc()) \
        .offset(offset).limit(per_page) \
        .all()

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        rows_html = render_template('_member_ids_rows.html', member_ids=member_ids)
        pagination_html = render_template('_pagination.html', page=page, total_pages=total_pages)
        return jsonify(rows=rows_html, pagination=pagination_html)

    return render_template(
        'member_ids.html',
        member_ids=member_ids,
        page=page,
        total_pages=total_pages,
        total=total
    )


@app.route('/add_member_ids', methods=['POST'])
@admin_required
def add_member_ids():
    try:
        start_num = int(request.form['start_num'])
        end_num = int(request.form['end_num'])

        # Tạo danh sách ID mới
        new_ids = [f"MEM-{str(i).zfill(3)}" for i in range(start_num, end_num + 1)]

        # Kiểm tra ID đã tồn tại
        existing_ids = set(
            db.session.query(MemberID.member_id)
            .filter(MemberID.member_id.in_(new_ids))
            .all()
        )
        existing_ids = {x[0] for x in existing_ids}

        to_insert = [MemberID(member_id=mid) for mid in new_ids if mid not in existing_ids]
        added_count = len(to_insert)

        if added_count > 0:
            db.session.bulk_save_objects(to_insert)
            db.session.commit()

            log_activity(
                "Thêm mã thành viên",
                f"{current_user.display_name} đã thêm {added_count} mã từ {new_ids[0]} đến {new_ids[-1]}."
            )

            return jsonify(success=True, added=added_count)
        else:
            return jsonify(success=False, message="Không có mã mới được thêm.")

    except Exception as e:
        db.session.rollback()
        return jsonify(success=False, message=str(e)), 500
    
# Tạo bảng mã thành viên
@app.route('/member_ids_table')
@admin_required
def member_ids_table():
    member_ids = db.session.query(MemberID).order_by(MemberID.created_at.desc()).all()
    return render_template('_member_ids_table.html', member_ids=member_ids)

@app.route('/delete_member_ids', methods=['POST'])
@admin_required
def delete_member_ids():
    start_id = request.form['start_id'].strip()
    end_id = request.form['end_id'].strip()

    try:
        start_num = int(start_id.split('-')[1])
        end_num = int(end_id.split('-')[1])
    except (IndexError, ValueError):
        return jsonify(success=False, message='Định dạng mã không hợp lệ!'), 400

    if start_num > end_num:
        return jsonify(success=False, message='Mã bắt đầu phải nhỏ hơn hoặc bằng mã kết thúc!'), 400

    ids_to_delete = [f"MEM-{str(i).zfill(3)}" for i in range(start_num, end_num + 1)]
    deleted = MemberID.query.filter(
        MemberID.member_id.in_(ids_to_delete),
        MemberID.is_used == False
    ).delete(synchronize_session=False)

    db.session.commit()

    log_activity(
        "Xoá mã thành viên",
        f"{current_user.display_name} đã xoá {deleted} mã thành viên chưa sử dụng (từ {start_id} đến {end_id})."
    )

    return jsonify(success=True, deleted=deleted)


@app.route('/update_points/<int:member_id>', methods=['POST'])
@admin_required
def update_points(member_id):
    points_change = int(request.form['points_change'])
    reason = request.form['reason']

    user = User.query.get(member_id)
    if user:
        # Nếu không cho tự cộng điểm, bật đoạn này
        # if user.id == session['user_id']:
        #     flash('Bạn không thể tự cộng điểm cho chính mình.', 'warning')
        #     return redirect(request.referrer or url_for('dashboard'))

        user.points += points_change
        log = PointLog(member_id=member_id,
                       points_change=points_change,
                       reason=reason,
                       admin_id=session['user_id'])
        db.session.add(log)
        db.session.commit()
        log_activity("Thay đổi điểm", f"{current_user.username}: cập nhật {points_change:+} điểm cho {user.username} (ID {user.id}) — lý do: {reason}")
        flash('Cập nhật điểm thành công!', 'success')
    else:
        flash('Không tìm thấy người dùng.', 'danger')

    return redirect(request.referrer or url_for('dashboard'))

@app.route('/logout')
def logout():
    session.clear()
    flash('Đã đăng xuất thành công.', 'success')
    return redirect(url_for('login'))

@app.route('/delete_member/<int:user_id>', methods=['POST'])
@admin_required
def delete_member(user_id):
    user = User.query.get(user_id) if user_id else None
    if user:
        member_id_obj = MemberID.query.filter_by(member_id=user.member_id).first()
        if member_id_obj:
            member_id_obj.is_used = False
            member_id_obj.used_by = None

        db.session.delete(user)
        db.session.commit()
        log_activity("Xoá thành viên", f"{current_user.username} đã xoá thành viên {user.username} (ID {user.id}, mã {user.member_id}).")
        flash('Đã xóa thành viên và giải phóng mã thành viên.', 'success')
    else:
        flash('Không tìm thấy người dùng.', 'error')

    return redirect(url_for('members'))

@app.route('/register_admin', methods=['GET', 'POST'])
@admin_required
def register_admin():
    if request.method == 'POST':
        member_id = request.form['member_id']
        display_name = request.form['display_name']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        if password != confirm_password:
            flash('Mật khẩu xác nhận không khớp.', 'error')
            return render_template('register_admin.html')

        existing_user = User.query.filter_by(member_id=member_id).first()
        if existing_user:
            flash('Mã admin đã tồn tại.', 'error')
            return render_template('register_admin.html')

        password_hash = generate_password_hash(password)
        new_admin = User(member_id=member_id, display_name=display_name,
                         password_hash=password_hash, role='admin', points=10)
        db.session.add(new_admin)
        db.session.commit()

        flash('Tạo tài khoản admin thành công!', 'success')
        return redirect(url_for('login'))

    return render_template('register_admin.html')

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    user = User.query.get(session['user_id'])

    if request.method == 'POST':
        display_name = request.form['display_name']
        password_current = request.form.get('password_current')
        password_new = request.form.get('password_new')
        password_confirm = request.form.get('password_confirm')

        user.display_name = display_name

        if password_current and password_new and password_confirm:
            if not check_password_hash(user.password_hash, password_current):
                flash('Mật khẩu hiện tại không đúng.', 'error')
            elif password_new != password_confirm:
                flash('Mật khẩu mới và xác nhận không khớp.', 'error')
            else:
                user.password_hash = generate_password_hash(password_new)
                flash('Đổi mật khẩu thành công.', 'success')
        else:
            flash('Cập nhật thông tin thành công.', 'success')

        db.session.commit()
        return redirect(url_for('profile'))

    return render_template('profile.html', user=user)

@app.route('/admins')
@login_required
@admin_required
def admins():
    per_page = 30
    page = int(request.args.get('page', 1))
    offset = (page - 1) * per_page

    total_admins = User.query.filter_by(role='admin').count()
    total_pages = ceil(total_admins / per_page)

    admins = User.query.filter_by(role='admin') \
        .order_by(User.created_at.desc()) \
        .offset(offset).limit(per_page).all()

    members = User.query.filter_by(role='member').all()

    can_create = current_user.role == 'admin'
    can_edit = current_user.member_id == 'ADMIN-001'

    return render_template('admins.html', admins=admins, members=members,
                       can_create=can_create, can_edit=can_edit,
                       page=page, total_pages=total_pages, total=total_admins)


@app.route('/delete_admin/<int:user_id>', methods=['POST'])
@admin_required
def delete_admin(user_id):
    if not current_user or current_user.member_id != 'ADMIN-001':
        flash('Bạn không có quyền xóa admin.', 'danger')
        return redirect(url_for('admins'))

    if user_id == current_user.id:
        log_activity("Xoá admin thất bại", f"{current_user.username} cố xoá chính mình (ID {current_user.id}) — bị từ chối.")
        flash('Không thể tự xóa chính mình.', 'danger')
        return redirect(url_for('admins'))

    admin = User.query.get(user_id)
    if admin and admin.role == 'admin':
        db.session.delete(admin)
        db.session.commit()
        log_activity("Xoá admin", f"{current_user.username} đã xoá admin {admin.username} (ID {admin.id}).")
        flash('Đã xóa admin thành công.', 'success')
    else:
        flash('Không tìm thấy admin.', 'error')

    return redirect(url_for('admins'))

@app.route('/update_admin_points/<int:user_id>', methods=['POST'])
@admin_required
def update_admin_points(user_id):


    if not current_user or current_user.member_id != 'ADMIN-001':
        flash('Bạn không có quyền cập nhật điểm admin.', 'danger')
        return redirect(url_for('admins'))

    admin = User.query.get(user_id)
    if admin and admin.role == 'admin':
        try:
            points = int(request.form['points'])
            admin.points = points
            db.session.commit()
            log_activity("Cập nhật điểm admin", f"{current_user.username} đã cập nhật điểm cho {admin.username} (ID {admin.id}) thành {points} điểm.")
            flash('Cập nhật điểm thành công.', 'success')
        except ValueError:
            flash('Giá trị điểm không hợp lệ.', 'danger')
    else:
        flash('Không tìm thấy admin.', 'danger')

    return redirect(url_for('admins'))

from sqlalchemy.inspection import inspect

@app.route('/download_db')
@admin_required
def download_db():
    user = User.query.get(session['user_id'])
    if not user or user.member_id not in ['ADMIN-001', 'ADMIN-030']:
        flash('Bạn không có quyền tải xuống cơ sở dữ liệu.', 'error')
        return redirect(url_for('dashboard'))

    temp_dir = tempfile.TemporaryDirectory()
    zip_path = os.path.join(temp_dir.name, "full_database_export.zip")

    def add_csv_to_zip(zipf, filename, headers, rows):
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(headers)
        writer.writerows(rows)
        zipf.writestr(filename, buffer.getvalue())
        buffer.close()

    with ZipFile(zip_path, 'w') as zipf:
        for mapper in db.Model.registry.mappers:
            model_class = mapper.class_
            if not hasattr(model_class, '__tablename__'):
                continue

            table_name = model_class.__tablename__ + '.csv'
            instances = model_class.query.all()
            if not instances:
                continue

            columns = [column.key for column in inspect(model_class).columns]
            rows = []
            for instance in instances:
                row = []
                for col in columns:
                    val = getattr(instance, col)
                    if isinstance(val, datetime):
                        val = val.strftime('%Y-%m-%d %H:%M:%S')
                    row.append(val)
                rows.append(row)

            add_csv_to_zip(zipf, table_name, columns, rows)

    return send_file(zip_path, as_attachment=True, download_name="full_database_export.zip")

# Luật sử dụng
@app.route('/rules', methods=['GET', 'POST'])
@admin_required
def rules():
    rule = Rule.query.first()
    if request.method == 'POST':
        content = request.form['content']
        if rule:
            rule.content = content
        else:
            rule = Rule(content=content)
            db.session.add(rule)
        db.session.commit()
        log_activity("Cập nhật luật", f"{current_user.username} đã cập nhật luật.")
        flash('Cập nhật nội dung luật thành công.', 'success')
        return redirect(url_for('rules'))
    return render_template('rules.html', rule=rule)

# Public view trên trang login
@cache.cached(timeout=300)
@app.route('/public_rules')
def public_rules():
    rule = Rule.query.first()
    return render_template('public_rules.html', rule=rule)

@app.route('/export_rules')
@admin_required
def export_rules():
    rule = Rule.query.first()
    if not rule:
        flash("Chưa có nội dung luật để xuất.", "warning")
        return redirect(url_for('rules'))

    doc = Document()
    doc.add_heading('Nội dung Luật', level=1)
    doc.add_paragraph(rule.content)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, as_attachment=True, download_name="luat.docx")

# Character abilities
FACTIONS = ["Phe Dân", "Phe Sói", "Phe Ba", "Đổi Phe"]

@cache.cached(timeout=300)
@app.route('/abilities')
@login_required
def abilities():
    # Lọc
    search_faction = request.args.get('faction', '').strip()
    search_name = request.args.get('name', '').strip()
    search_desc = request.args.get('desc', '').strip()

    query = CharacterAbility.query
    if search_faction:
        query = query.filter(CharacterAbility.faction.ilike(f"%{search_faction}%"))
    if search_name:
        query = query.filter(CharacterAbility.name.ilike(f"%{search_name}%"))
    if search_desc:
        query = query.filter(CharacterAbility.description.ilike(f"%{search_desc}%"))

    abilities = query.order_by(CharacterAbility.faction, CharacterAbility.order_in_faction).all()
    is_admin = session.get('user_role') == 'admin'

    # Tính STT kế tiếp cho từng phe
    FACTIONS = ['Phe Dân', 'Phe Sói', 'Phe Ba', 'Đổi Phe']
    next_orders = {
        faction: (db.session.query(db.func.max(CharacterAbility.order_in_faction))
                  .filter_by(faction=faction).scalar() or 0) + 1
        for faction in FACTIONS
    }

    # Nhóm chức năng theo phe và sắp xếp theo thứ tự FACTIONS
    grouped_abilities = {f: [] for f in FACTIONS}
    for a in abilities:
        grouped_abilities.setdefault(a.faction, []).append(a)

    # Đảm bảo thứ tự phe khi truyền vào template
    grouped_abilities_ordered = {f: grouped_abilities.get(f, []) for f in FACTIONS}

    return render_template(
        'abilities.html',
        abilities=abilities,
        is_admin=is_admin,
        next_orders=next_orders,
        grouped_abilities=grouped_abilities_ordered
    )


@app.route('/abilities/add', methods=['POST'])
@admin_required
def add_ability():
    faction = request.form['faction']
    order_in_faction = request.form.get('order')

    if not order_in_faction or order_in_faction.strip() == '':
        max_order = db.session.query(db.func.max(CharacterAbility.order_in_faction)).filter_by(faction=faction).scalar()
        order_in_faction = (max_order or 0) + 1
    else:
        order_in_faction = int(order_in_faction)

    name = request.form['name']
    description = request.form['description']

    new_ability = CharacterAbility(
        faction=faction,
        order_in_faction=order_in_faction,
        name=name,
        description=description
    )
    db.session.add(new_ability)
    db.session.commit()
    cache.delete_memoized(abilities)
    log_activity("Thêm chức năng", f"{current_user.username} đã thêm chức năng mới: '{new_ability.name}' (ID {new_ability.id}) vào phe {new_ability.faction}.")
    flash('Đã thêm chức năng.', 'success')
    return redirect(url_for('abilities'))


@app.route('/abilities/edit/<int:ability_id>', methods=['POST'])
@admin_required
def edit_ability(ability_id):
    ability = CharacterAbility.query.get(ability_id)
    if ability:
        ability.faction = request.form['faction']
        ability.order_in_faction = int(request.form['order'])
        ability.name = request.form['name']
        ability.description = request.form['description']
        db.session.commit()
        cache.delete_memoized(abilities)
        log_activity("Sửa chức năng", f"{current_user.username} đã cập nhật chức năng '{ability.name}' (ID {ability.id}).")
        flash('Đã cập nhật.', 'success')
    else:
        flash('Không tìm thấy chức năng.', 'danger')
    return redirect(url_for('abilities'))


@app.route('/abilities/delete/<int:ability_id>', methods=['POST'])
@admin_required
def delete_ability(ability_id):
    ability = CharacterAbility.query.get(ability_id)
    if ability:
        db.session.delete(ability)
        db.session.commit()
        cache.delete_memoized(abilities)
        log_activity("Xóa chức năng", f"{current_user.username} đã xóa chức năng '{ability.name}' (ID {ability.id}).")
        flash('Đã xóa chức năng.', 'success')
    else:
        flash('Không tìm thấy chức năng.', 'danger')
    cache.delete_memoized(abilities)
    return redirect(url_for('abilities'))

# Kim Bài Miễn Tử
from math import ceil
from sqlalchemy import func, case

@app.route('/kim_bai')
@login_required
def kim_bai():
    per_page = 30
    page = int(request.args.get('page', 1))

    # 🔁 Gộp các count lại 1 truy vấn duy nhất
    counts = db.session.query(
        func.count(User.id).label('total'),
        func.sum(case((User.has_kim_bai == True, 1), else_=0)).label('has_kim_bai_count'),
        func.sum(case((User.has_kim_bai == False, 1), else_=0)).label('no_kim_bai_count')
    ).first()

    total = counts.total or 0
    has_kim_bai_count = counts.has_kim_bai_count or 0
    no_kim_bai_count = counts.no_kim_bai_count or 0

    # 📄 Truy vấn danh sách người dùng (có phân trang)
    members = User.query.order_by(User.display_name) \
        .offset((page - 1) * per_page).limit(per_page).all()

    total_pages = ceil(total / per_page)

    return render_template(
        'kim_bai.html',
        members=members,
        total=total,
        page=page,
        total_pages=total_pages,
        has_kim_bai_count=has_kim_bai_count,
        no_kim_bai_count=no_kim_bai_count
    )

@app.route('/increase_death/<int:user_id>', methods=['POST'])
@admin_required
def increase_death(user_id):
    user = User.query.get(user_id) if user_id else None
    if user:
        user.death_count += 1
        # Nếu death_count chia hết cho 2 → cấp kim bài
        if user.death_count > 0 and user.death_count % 2 == 0:
            user.has_kim_bai = True
        db.session.commit()

        log = KimBaiLog(user_id=user.id, timestamp=datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')))
        db.session.add(log)
        db.session.commit()
        cache.delete_memoized(top_tier)
        log_activity("Tăng lượt chết", f"{current_user.username} tăng lượt chết cho {user.display_name} (ID {user.id}). Tổng: {user.death_count}.")
        flash('Đã tăng lượt chết.', 'success')
    return redirect(url_for('kim_bai'))

@app.route('/use_kim_bai/<int:user_id>', methods=['POST'])
@admin_required
def use_kim_bai(user_id):
    user = User.query.get(user_id) if user_id else None
    if user and user.has_kim_bai:
        user.has_kim_bai = False
        db.session.commit()
        log_activity("Sử dụng kim bài", f"{current_user.username} đánh dấu {user.display_name} (ID {user.id}) đã sử dụng kim bài.")
        flash('Đã sử dụng kim bài.', 'success')
    else:
        flash('Không có kim bài để dùng.', 'danger')
    return redirect(url_for('kim_bai'))

@app.route('/decrease_death/<int:user_id>', methods=['POST'])
@admin_required
def decrease_death(user_id):
    user = User.query.get(user_id) if user_id else None
    if user and user.death_count > 0:
        # Trừ lượt chết
        user.death_count -= 1

        # Xóa dòng log gần nhất của người này
        last_log = KimBaiLog.query.filter_by(user_id=user.id)\
                                  .order_by(KimBaiLog.timestamp.desc())\
                                  .first()
        if last_log:
            db.session.delete(last_log)

        # Cập nhật kim bài
        if user.death_count > 0 and user.death_count % 2 == 0:
            user.has_kim_bai = True
        else:
            user.has_kim_bai = False

        db.session.commit()
        cache.delete_memoized(top_tier)
        log_activity("Giảm lượt chết", f"{current_user.username} giảm lượt chết cho {user.display_name} (ID {user.id}). Còn {user.death_count} lượt chết.")
        flash('Đã giảm lượt chết.', 'success')
    else:
        flash('Không thể giảm nữa.', 'warning')
    return redirect(url_for('kim_bai'))

#Top
@cache.cached(timeout=300)
@app.route('/top_tier')
@login_required
def top_tier():
    # Top 3 chết nhiều nhất tháng
    vietnam_now = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))
    current_month = vietnam_now.month
    current_year = vietnam_now.year

    top_deaths = (
        db.session.query(User.display_name, func.count(KimBaiLog.id).label("death_count"))
        .join(KimBaiLog)
        .filter(func.extract('month', KimBaiLog.timestamp) == current_month)
        .filter(func.extract('year', KimBaiLog.timestamp) == current_year)
        .group_by(User.id)
        .order_by(func.count(KimBaiLog.id).desc())
        .limit(3)
        .all()
    )

    return render_template("top_tier.html", top_deaths=top_deaths, now=datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')))


# Blacklist management
@cache.cached()
@app.route('/blacklist', methods=['GET', 'POST'])
@login_required
def blacklist():
    role_filter = request.args.get('role')
    user_filter_id = request.args.get('user_id')

    # Lọc danh sách blacklist
    query = BlacklistEntry.query

    if role_filter == 'admin':
        query = query.join(User).filter(User.role == 'admin')
    elif role_filter == 'member':
        query = query.join(User).filter(User.role == 'member')

    if user_filter_id:
        query = query.filter(BlacklistEntry.created_by_id == int(user_filter_id))
    
    per_page = 30
    page = int(request.args.get('page', 1))
    offset = (page - 1) * per_page

    total = query.count()
    total_pages = ceil(total / per_page)
    entries = query.order_by(BlacklistEntry.id.desc()).all()

    # Danh sách người đã từng tạo entry
    creators = (
        db.session.query(User)
        .join(BlacklistEntry, BlacklistEntry.created_by_id == User.id)
        .filter(User.role.in_(['admin', 'member']))
        .distinct()
        .all()
    )

    return render_template(
        'blacklist.html',
        entries=entries,
        all_users=creators,
        role_filter=role_filter,
        user_filter_id=user_filter_id,
        user=current_user,
        page=page,
        total_pages=total_pages,
        total=total
    )


@app.route('/add_blacklist', methods=['POST'])
@login_required
def add_blacklist():
    name = request.form['name']
    facebook_link = request.form.get('facebook_link')
    if not name.strip():
        flash('Tên là bắt buộc!', 'danger')
        return redirect(url_for('blacklist'))

    new_entry = BlacklistEntry(
        name=name.strip(),
        facebook_link=facebook_link.strip() if facebook_link else None,
        created_by_id=session['user_id']
    )
    db.session.add(new_entry)
    db.session.commit()
    cache.delete_memoized(blacklist)
    log_activity("Thêm blacklist", f"{current_user.username} thêm '{new_entry.name}' vào blacklist.")
    flash('Đã thêm vào blacklist.', 'success')
    return redirect(url_for('blacklist'))

@app.route('/delete_blacklist/<int:entry_id>', methods=['POST'])
@login_required
def delete_blacklist(entry_id):
    entry = BlacklistEntry.query.get(entry_id)


    if entry and (entry.created_by_id == current_user.id or current_user.member_id == 'ADMIN-001'):
        db.session.delete(entry)
        db.session.commit()
        cache.delete_memoized(blacklist)
        log_activity("Xóa blacklist", f"{current_user.username} đã xóa blacklist entry ID {entry.id}.")
        flash('Đã xoá mục khỏi blacklist.', 'success')
    else:
        flash('Bạn không có quyền xoá mục này.', 'danger')
    return redirect(url_for('blacklist'))

@app.route('/edit_blacklist_author/<int:entry_id>', methods=['POST'])
@admin_required
def edit_blacklist_author(entry_id):

    if current_user.member_id != 'ADMIN-001':
        flash('Bạn không có quyền sửa người nhập!', 'danger')
        return redirect(url_for('blacklist'))

    new_user_id = request.form.get('new_user_id')
    entry = BlacklistEntry.query.get(entry_id)
    user_target = User.query.get(new_user_id)

    if entry and user_target:
        entry.created_by_id = user_target.id
        db.session.commit()
        cache.delete_memoized(blacklist)
        log_activity("Cập nhật người nhập blacklist", f"{current_user.username} sửa created_by_id của entry ID {entry.id} thành user ID {user_target.id}.")
        flash('Đã cập nhật người nhập.', 'success')
    else:
        flash('Không tìm thấy người dùng hoặc mục!', 'danger')
    return redirect(url_for('blacklist'))

@app.route("/game_history")
def game_history():
    from models import GameHistory, User, CharacterAbility

    faction_order = {
        "Phe Dân": 1,
        "Phe Sói": 2,
        "Phe Ba": 3,
        "Đổi Phe": 4
    }

    games = GameHistory.query.order_by(GameHistory.created_at.desc()).all()
    chars = CharacterAbility.query.all()
    chars_sorted = sorted(chars, key=lambda c: faction_order.get(c.faction, 999))
    users = User.query.order_by(User.member_id.asc()).all()

    FACTION_ICONS = {
        "Phe Dân": ("fa-users", "bg-success text-white"),
        "Phe Sói": ("fa-brands fa-wolf-pack-battalion", "bg-danger-subtle text-danger"),
        "Phe Ba": ("fa-user-secret", "bg-secondary-subtle text-dark"),
        "Đổi Phe": ("fa-random", "bg-warning-subtle text-warning")
    }

    # ✅ Chuyển users và chars sang dict để dùng với tojson trong template
    user_dicts = [
        {"id": u.id, "display_name": u.display_name, "member_id": u.member_id}
        for u in users
    ]
    char_dicts = [
        {"id": c.id, "name": c.name, "faction": c.faction}
        for c in chars_sorted
    ]

    return render_template(
        "game_history.html",
        games=games,
        users=users,
        chars=chars_sorted,
        FACTION_ICONS=FACTION_ICONS,
        user_dicts=user_dicts,         # ✅ thêm
        char_dicts=char_dicts          # ✅ thêm
    )


import random
from flask import request, redirect, url_for, flash
from models import GameHistory, GamePlayer, User, PointLog, db

@app.route("/create_game", methods=["POST"])
@login_required
def create_game():
    from models import GameHistory, GamePlayer, User, PointLog, db

    mode = request.form.get("mode")

    # ✅ PHÂN THỦ CÔNG
    if mode == "manual":
        manual_players = request.form.getlist('manual_players[]')
        manual_chars = request.form.getlist('manual_chars[]')

        # print("👤 Người chơi:", manual_players)
        # print("🎭 Nhân vật:", manual_chars)

        if len(manual_players) != len(manual_chars) or len(manual_players) == 0:
            flash("Số lượng người chơi và nhân vật phải bằng nhau và lớn hơn 0.", "danger")
            return redirect(url_for('game_history'))

        new_game = GameHistory(host_id=session['user_id'])
        db.session.add(new_game)
        db.session.commit()

        for pid, cid in zip(manual_players, manual_chars):
            db.session.add(GamePlayer(game_id=new_game.id, player_id=pid, char_id=cid))
        db.session.commit()

        # ✅ Cộng điểm
        for pid in manual_players:
            user = User.query.get(pid)
            if user and user.points < 10:
                before = user.points
                user.points = min(user.points + 2, 10)
                db.session.add(PointLog(
                    member_id=user.id,
                    points_change=user.points - before,
                    reason="Tham gia ván chơi",
                    admin_id=session.get("user_id")
                ))
                log_activity("Cộng điểm", f"{current_user.username} cộng {user.points - before} điểm cho {user.display_name} (ID {user.id}) trong ván chơi.")
                # print(f"✔️ +{user.points - before} điểm cho {user.display_name} (ID {user.id}): {before} ➜ {user.points}")
        db.session.commit()
        log_activity("Tạo ván chơi", f"{current_user.username} tạo ván chơi (thủ công), game ID {new_game.id}, {len(manual_players)} người chơi.")
        flash("Đã tạo ván chơi phân thủ công!", "success")
        return redirect(url_for('game_history'))

    # ✅ PHÂN NGẪU NHIÊN
    elif mode == "random":
        player_ids = request.form.getlist("players")
        char_ids_str = request.form.get("char_ids", "")
        char_ids = [int(cid) for cid in char_ids_str.split(',') if cid.strip().isdigit()]

        # print("🧪 Form raw:", request.form)
        # print("👤 Người chơi:", player_ids)
        # print("🎭 Nhân vật:", char_ids)

        if len(player_ids) != len(char_ids) or len(player_ids) == 0:
            flash("Số lượng người chơi và nhân vật phải bằng nhau và lớn hơn 0.", "danger")
            return redirect(url_for('game_history'))

        new_game = GameHistory(host_id=session['user_id'])
        db.session.add(new_game)
        db.session.flush()

        for pid, cid in zip(player_ids, char_ids):
            db.session.add(GamePlayer(game_id=new_game.id, player_id=int(pid), char_id=int(cid)))
        db.session.commit()

        # ✅ Cộng điểm
        for pid in player_ids:
            user = User.query.get(pid)
            if user and user.points < 10:
                before = user.points
                user.points = min(user.points + 2, 10)
                db.session.add(PointLog(
                    member_id=user.id,
                    points_change=user.points - before,
                    reason="Tham gia ván chơi",
                    admin_id=session.get("user_id")
                ))
                log_activity("Cộng điểm", f"{current_user.username} cộng {user.points - before} điểm cho {user.display_name} (ID {user.id}) trong ván chơi.")
                # print(f"✔️ +{user.points - before} điểm cho {user.display_name} (ID {user.id}): {before} ➜ {user.points}")
        db.session.commit()
        log_activity("Tạo ván chơi", f"{current_user.username} tạo ván chơi (ngẫu nhiên), game ID {new_game.id}, {len(player_ids)} người chơi.")
        flash("Tạo ván (phân ngẫu nhiên) thành công!", "success")
        return redirect(url_for('game_history'))

    # ❌ Trường hợp không xác định
    flash("Dữ liệu không hợp lệ.", "danger")
    return redirect(url_for('game_history'))

@app.route('/update_game_note/<int:game_id>', methods=['POST'])
@admin_required
def update_game_note(game_id):
    game = GameHistory.query.get_or_404(game_id)
    print("Before:", game.notes, game.tags)

    game.notes = request.form.get('note', '')  # sửa lại đúng name
    selected_tags = request.form.getlist('tags')
    game.tags = ",".join(selected_tags)

    print("After:", game.notes, game.tags)

    db.session.commit()
    log_activity("Cập nhật ván chơi", f"Admin {current_user.username} cập nhật note và tag cho ván chơi ID {game.id}.")
    flash('Đã cập nhật ván chơi.', 'success')
    return redirect(url_for('game_history'))

@app.route('/delete_game/<int:game_id>')
@admin_required
def delete_game(game_id):
    game = GameHistory.query.get_or_404(game_id)
    db.session.delete(game)
    db.session.commit()
    log_activity("Xóa ván chơi", f"Admin {current_user.username} đã xóa ván chơi ID {game.id}.")
    flash('Đã xóa ván chơi.', 'success')
    return redirect(url_for('game_history'))

@app.route("/day_off", methods=["GET", "POST"])
@login_required
def day_off():
    from datetime import datetime
    from models import User, PlayerOffRequest

    user_id = session.get("user_id")
    user = User.query.get(user_id) if user_id else None

    if request.method == "POST":
        start_date = datetime.strptime(request.form["start_date"], "%Y-%m-%d").date()
        end_date = datetime.strptime(request.form["end_date"], "%Y-%m-%d").date()
        reason = request.form.get("reason", "")

        actual_user_id = user.id
        created_by = user.id

        # ✅ Nếu là admin thì được chọn người khác để tạo hộ
        if user.role == 'admin' and request.form.get("user_id"):
            actual_user_id = int(request.form["user_id"])

        request_off = PlayerOffRequest(
            user_id=actual_user_id,
            start_date=start_date,
            end_date=end_date,
            reason=reason,
            created_by=created_by
        )
        db.session.add(request_off)
        db.session.commit()
        if actual_user_id == user.id:
            log_activity("Gửi yêu cầu nghỉ", f"{user.username} xin nghỉ từ {start_date} đến {end_date}.")
        else:
            log_activity("Tạo yêu cầu nghỉ hộ", f"Admin {user.username} tạo yêu cầu nghỉ từ {start_date} đến {end_date} cho user ID {actual_user_id}.")

        flash("✔️ Đã gửi yêu cầu xin nghỉ!", "success")
        return redirect(url_for("day_off"))

    # ✅ Tất cả người dùng đều thấy danh sách đầy đủ
    offs = PlayerOffRequest.query.order_by(PlayerOffRequest.start_date.desc()).all()

    # ✅ Nếu là admin, cung cấp danh sách users để chọn trong form
    users = User.query.order_by(User.member_id.asc()).all() if user.role == 'admin' else []

    return render_template(
        "day_off.html",
        offs=offs,
        users=users,
        current_user=user,
        user=user  # dùng cho template nếu cần
    )

@app.route("/delete_off/<int:off_id>", methods=["POST"])
@login_required
def delete_off(off_id):
    user_id = session.get("user_id")
    user = User.query.get(user_id) if user_id else None

    if user.role != 'admin':
        flash("Bạn không có quyền xóa yêu cầu nghỉ!", "danger")
        return redirect(url_for("day_off"))

    off = PlayerOffRequest.query.get_or_404(off_id)
    db.session.delete(off)
    db.session.commit()
    log_activity("Xóa yêu cầu nghỉ", f"Admin {user.username} đã xóa yêu cầu nghỉ ID {off.id} của user ID {off.user_id}.")
    flash("✔️ Đã xóa yêu cầu nghỉ!", "success")
    return redirect(url_for("day_off"))


from datetime import datetime, timedelta
from sqlalchemy import func
from flask import render_template

from sqlalchemy import func, union_all, select
from sqlalchemy.orm import lazyload, joinedload

@app.route("/frequency")
@login_required
def frequency():
    vietnam_now = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))
    today = vietnam_now.date()

    # 🔎 Tải tất cả user và map theo ID để tra nhanh
    users = User.query.options(lazyload("*")).filter_by(role='member').all()
    user_map = {u.id: u for u in users}

    # 📌 Lấy danh sách nghỉ (vẫn còn hiệu lực)
    current_offs = db.session.query(
        PlayerOffRequest.user_id,
        func.max(PlayerOffRequest.end_date).label("latest_end")
    ).filter(
        PlayerOffRequest.start_date <= today,
        PlayerOffRequest.end_date >= today
    ).group_by(PlayerOffRequest.user_id).all()

    off_dict = {user_id: latest_end for user_id, latest_end in current_offs}

    # 📌 Union subquery giữa người chơi và host
    gp_sub = db.session.query(
        GamePlayer.player_id.label("user_id"),
        GameHistory.created_at.label("played_at")
    ).join(GameHistory, GamePlayer.game_id == GameHistory.id)

    host_sub = db.session.query(
        GameHistory.host_id.label("user_id"),
        GameHistory.created_at.label("played_at")
    )

    union_q = gp_sub.union_all(host_sub).subquery()

    # 📌 Tính số lượt chơi và lần chơi gần nhất
    stats = db.session.query(
        union_q.c.user_id,
        func.count().label("play_count"),
        func.max(union_q.c.played_at).label("last_play")
    ).group_by(union_q.c.user_id).all()

    data = []
    handled_ids = set()

    for user_id, play_count, last_play in stats:
        user = user_map.get(user_id)
        if not user:
            continue

        if user_id in off_dict:
            adjusted_last_play = max(last_play.date(), off_dict[user_id])
        else:
            adjusted_last_play = last_play.date()

        inactive = (today - adjusted_last_play).days > 7

        data.append({
            "user": user,
            "play_count": play_count,
            "last_play": adjusted_last_play,
            "inactive": inactive,
            "on_leave": user_id in off_dict
        })
        handled_ids.add(user_id)

    # 📌 Những người chưa từng chơi
    for u in users:
        if u.id in handled_ids:
            continue

        if u.id in off_dict:
            adjusted_last_play = off_dict[u.id]
            inactive = (today - adjusted_last_play).days > 7
            data.append({
                "user": u,
                "play_count": 0,
                "last_play": adjusted_last_play,
                "inactive": inactive,
                "on_leave": True
            })
        else:
            data.append({
                "user": u,
                "play_count": 0,
                "last_play": None,
                "inactive": True,
                "on_leave": False
            })

    per_page = 30
    page = int(request.args.get('page', 1))
    offset = (page - 1) * per_page
    total = len(data)
    total_pages = ceil(total / per_page)

    data = data[offset:offset + per_page]

    return render_template("frequency.html", data=data, page=page, total=total, total_pages=total_pages)

@app.route('/activity_log')
@admin_required
def activity_log():
    logs = cache.get('active_logs_vn')
    if logs is None:
        raw_logs = ActivityLog.query.order_by(ActivityLog.timestamp.desc()).all()
        logs = convert_logs_timezone(raw_logs)
        cache.set('active_logs_vn', logs, timeout=300)
    return render_template("activity_log.html", logs=logs)

# Danh sách preset để hiện mô tả cho từng theme
THEME_PRESETS = {
    'default': {
        'name': 'Mặc định',
        'description': 'Theme cơ bản với màu xanh tím',
        'effects': False
    },
    'dark': {
        'name': 'Tối',
        'description': 'Theme tối như cái đầu của bạn',
        'effects': False
    },
    'sakura': {
        'name': 'Hoa anh đào',
        'description': 'Theme hồng nhẹ nhàng, hồng mạnh mẽ, hồng nam tính',
        'effects': True
    },
    'galaxy': {
        'name': 'Thiên hà',
        'description': 'Theme vũ trụ',
        'effects': True
    },
    'ocean': {
        'name': 'Đại dương',
        'description': 'Theme xanh dương như chờ lương mỗi phút',
        'effects': True
    },
    'forest': {
        'name': 'Rừng xanh',
        'description': 'Theme xanh như greenflag mà em chê',
        'effects': True
    },
    'sunset': {
        'name': 'Hoàng hôn',
        'description': 'Theme cam vàng màu nắng',
        'effects': True
    },
    'halloween': {
        'name': 'Halloween',
        'description': 'Theme Halloween 🎃',
        'effects': True
    },
    'christmas': {
        'name': 'Giáng Sinh',
        'description': 'Theme tuyết rơi & màu Noel 🎄',
        'effects': True
    },
    'newyear': {
        'name': 'Năm Mới',
        'description': 'Theme chúc mừng năm mới 🎉',
        'effects': True
    },
}

# Trả về theme hiệu lực (ưu tiên ngày lễ)
def get_theme(user):
    today = datetime.now(pytz.timezone('Asia/Ho_Chi_Minh'))

    # Ưu tiên theme đặc biệt theo ngày lễ
    if today.month == 10 and today.day >= 25:
        return 'halloween'
    elif today.month == 12 and today.day >= 24:
        return 'christmas'
    elif today.month == 1 and today.day <= 2:
        return 'newyear'

    # Nếu user có theme cá nhân thì dùng
    return user.theme if user and user.theme else 'default'

@app.route('/change-theme', methods=['GET', 'POST'])
def change_theme():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    themes = list(THEME_PRESETS.keys())

    if request.method == 'POST':
        selected = request.form.get('theme')
        if selected in themes:
            user.theme = selected
            db.session.commit()

            # ✅ Cập nhật session
            cache.set(f"user_theme:{user.id}", selected, timeout=3600)

            flash(f'Đã đổi giao diện sang theme: {selected}', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Theme không hợp lệ.', 'danger')

    return render_template('change_theme.html', user=user, themes=themes, THEME_PRESETS=THEME_PRESETS)


import flask
from datetime import datetime

@app.route("/version")
@login_required
def show_version():
    user = User.query.get(session['user_id'])
    if not user or user.role != 'admin':
        flash("Bạn không có quyền truy cập trang này.", "danger")
        return redirect(url_for('dashboard'))

    # Múi giờ Việt Nam
    vietnam_tz = pytz.timezone('Asia/Ho_Chi_Minh')
    
    recent_logs = ActivityLog.query.filter(ActivityLog.action == "Nâng cấp hệ thống") \
                                   .order_by(ActivityLog.timestamp.desc()) \
                                   .limit(5).all()
    
    # Chuyển đổi timestamp sang múi giờ Việt Nam
    for log in recent_logs:
        if log.timestamp.tzinfo is None:
            # Nếu timestamp không có timezone, giả định là UTC
            log.timestamp = pytz.utc.localize(log.timestamp)
        log.timestamp = log.timestamp.astimezone(vietnam_tz)
    
    # Trả về JSON data cho AJAX request
    if request.headers.get('Content-Type') == 'application/json' or request.args.get('ajax'):
        return jsonify({
            'version': APP_VERSION,
            'flask_version': flask.__version__,
            'changelog': APP_CHANGELOG,
            'logs': [{
                'timestamp': log.timestamp.strftime('%d-%m-%Y %H:%M'),
                'detail': log.detail
            } for log in recent_logs],
            'release_date': '22/07/2025',  # Hoặc lấy từ database
            'build_number': f"#{APP_VERSION}.{datetime.now(vietnam_tz).strftime('%Y%m%d')}"
        })
    
    # Nếu không phải AJAX request, redirect về dashboard với modal trigger
    return redirect(url_for('dashboard', show_version='true'))

VERSION_FILE = os.path.join(os.path.dirname(__file__), 'version.txt')
if os.path.exists(VERSION_FILE):
    with open(VERSION_FILE) as f:
        APP_VERSION = f.read().strip()
else:
    APP_VERSION = "v0.0.0"